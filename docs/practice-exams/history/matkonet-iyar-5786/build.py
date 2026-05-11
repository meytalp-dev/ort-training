# -*- coding: utf-8 -*-
"""
Build matkonet (practice exam) for Tolodot Am Yisrael — שאלון 30281, אייר תשפ"ו.
Based on: קיץ תשפ"ג, with question replacements from תשפ"ה, תשפ"ד, תשפ"ב.

Replacements (base = תשפ"ג):
  Q2 -> Q1 from תשפ"ה     (גזרת השחיטה)
  Q3 -> Q5 from תשפ"ה     (היציאה מחוץ לחומות)
  Q4 -> Q6 from תשפ"ד     (מאורעות תרפ"ט)
  Q5 -> Q5 from תשפ"ד     (ארץ-ישראל במלחה"ע הראשונה)
  Q9 -> Q7 from תשפ"ה     (הבעש"ט)
  Q10 -> Q7 from תשפ"ד    (4 sub-items, student picks one)
  Q11 -> Q12 from תשפ"ב   (משפט אייכמן)

Structure:
  Perek I   (29pts, pick 1/2)
  Perek II  (50pts, pick 2/4, 25pts each)
  Perek III (21pts, pick 3/5, 7pts each)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

OUT_DOCX = r"c:\Users\meyta\Downloads\ort-presentation-builder\docs\practice-exams\history\matkonet-iyar-5786\matkonet-iyar-5786.docx"

# ===================== RTL helpers =====================

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_p(doc, text="", *, bold=False, underline=False, size=11,
          align='right', font='David', space_before=0, space_after=0,
          first_indent=None, indent_left=None, indent_right=None):
    p = doc.add_paragraph()
    set_rtl(p)
    pf = p.paragraph_format
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if indent_left is not None:
        pf.left_indent = Cm(indent_left)
    if indent_right is not None:
        pf.right_indent = Cm(indent_right)
    if text:
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.underline = underline
        # east asian and cs fonts too
        rpr = r._r.get_or_add_rPr()
        rFonts = rpr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rpr.insert(0, rFonts)
        rFonts.set(qn('w:cs'), font)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        # cs size
        sz = OxmlElement('w:szCs')
        sz.set(qn('w:val'), str(int(size * 2)))
        rpr.append(sz)
        # bold cs
        if bold:
            bCs = OxmlElement('w:bCs')
            rpr.append(bCs)
        set_run_rtl(r)
    return p

def add_runs(doc, runs, *, align='right', space_before=0, space_after=0,
             indent_right=None, indent_left=None, first_indent=None, hanging=None):
    """runs = list of dicts with keys: text, bold, underline, size, font"""
    p = doc.add_paragraph()
    set_rtl(p)
    pf = p.paragraph_format
    if align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if indent_right is not None:
        pf.right_indent = Cm(indent_right)
    if indent_left is not None:
        pf.left_indent = Cm(indent_left)
    if first_indent is not None:
        pf.first_line_indent = Cm(first_indent)
    if hanging is not None:
        pf.first_line_indent = Cm(-hanging)
        pf.left_indent = Cm(hanging) if pf.left_indent is None else pf.left_indent
    for rd in runs:
        text = rd.get('text', '')
        bold = rd.get('bold', False)
        underline = rd.get('underline', False)
        size = rd.get('size', 11)
        font = rd.get('font', 'David')
        r = p.add_run(text)
        r.font.name = font
        r.font.size = Pt(size)
        r.bold = bold
        r.underline = underline
        rpr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), font)
        rFonts.set(qn('w:ascii'), font)
        rFonts.set(qn('w:hAnsi'), font)
        rpr.insert(0, rFonts)
        sz = OxmlElement('w:szCs')
        sz.set(qn('w:val'), str(int(size * 2)))
        rpr.append(sz)
        if bold:
            bCs = OxmlElement('w:bCs')
            rpr.append(bCs)
        set_run_rtl(r)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

def add_horizontal_continue(doc):
    add_p(doc, "/המשך מעבר לדף/", align='left', size=10, space_before=6)

# ===================== Document setup =====================

doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'David'
style.font.size = Pt(11)

# Section: A4, bidi
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

sectPr = section._sectPr
bidi = OxmlElement('w:bidi')
sectPr.append(bidi)

# ===================== Cover page =====================

# Header table: right side = exam metadata, left side = State of Israel
tbl = doc.add_table(rows=1, cols=2)
tbl.autofit = False
tbl.columns[0].width = Cm(8)
tbl.columns[1].width = Cm(8)

# RIGHT cell (appears on RIGHT in RTL) = exam metadata — but in LTR table order cells[0] is left visually in LTR; since we're bidi, cells[0] is first-in-RTL = right
right_cell = tbl.cell(0, 0)
right_cell.text = ''
rp = right_cell.paragraphs[0]
set_rtl(rp)
rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def cell_line(cell, label, value, first=False):
    if first:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    for txt, bold in [(f"{label}:\t", False), (value, False)]:
        r = p.add_run(txt)
        r.font.name = 'David'
        r.font.size = Pt(11)
        rpr = r._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:cs'), 'David')
        rpr.insert(0, rFonts)
        set_run_rtl(r)

cell_line(right_cell, "סוג הבחינה", "בגרות (מתכונת)", first=True)
cell_line(right_cell, "מועד הבחינה", 'אייר תשפ"ו')
cell_line(right_cell, "מספר השאלון", "30281")

left_cell = tbl.cell(0, 1)
lp = left_cell.paragraphs[0]
set_rtl(lp)
lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = lp.add_run("מדינת ישראל")
r.font.name = 'David'
r.font.size = Pt(14)
r.bold = True
rpr = r._r.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:cs'), 'David'); rpr.insert(0, rFonts)
bCs = OxmlElement('w:bCs'); rpr.append(bCs)
set_run_rtl(r)

lp2 = left_cell.add_paragraph()
set_rtl(lp2)
lp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = lp2.add_run("משרד החינוך")
r.font.name = 'David'
r.font.size = Pt(12)
r.bold = True
rpr = r._r.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:cs'), 'David'); rpr.insert(0, rFonts)
bCs = OxmlElement('w:bCs'); rpr.append(bCs)
set_run_rtl(r)

# Remove table borders
for row in tbl.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

add_p(doc, "", space_before=18)
add_p(doc, "תולדות עם ישראל", bold=True, size=24, align='center', space_before=12, space_after=6)
add_p(doc, "לבתי ספר עצמאיים", bold=True, size=14, align='center', space_after=18)

add_p(doc, "הוראות", bold=True, size=14, align='center', space_after=12)

# א. משך הבחינה
add_runs(doc, [
    {'text': 'א.\t', 'size': 11},
    {'text': 'משך הבחינה', 'underline': True, 'size': 11},
    {'text': ': שעתיים וחצי.', 'size': 11},
], space_after=6, indent_right=0.5)

# ב. מבנה השאלון
add_runs(doc, [
    {'text': 'ב.\t', 'size': 11},
    {'text': 'מבנה השאלון ומפתח ההערכה', 'underline': True, 'size': 11},
    {'text': ': בשאלון זה שלושה פרקים.', 'size': 11},
], space_after=6, indent_right=0.5)

structure_lines = [
    ("פרק ראשון – שאלות תלויות קטע", "– (1×29) – 29 נקודות"),
    ("פרק שני – שאלות נושאים – סוגיות נבחרות בתולדות עם ישראל", "– (2×25) – 50 נקודות"),
    ("פרק שלישי – שאלות קצרות על מושגים מתוכנית הלימודים כולה", "– (3×7) – 21 נקודות"),
]
for label, pts in structure_lines:
    add_runs(doc, [
        {'text': label + '\t\t', 'size': 11},
        {'text': pts, 'size': 11},
    ], indent_right=1.5, space_after=4)

add_runs(doc, [
    {'text': '\t\t\t\t', 'size': 11},
    {'text': 'סך הכול 100 נקודות', 'size': 11, 'underline': True},
], indent_right=1.5, space_after=10)

add_runs(doc, [
    {'text': 'ג.\t', 'size': 11},
    {'text': 'חומר עזר מותר בשימוש', 'underline': True, 'size': 11},
    {'text': ': אין.', 'size': 11},
], space_after=4, indent_right=0.5)

add_runs(doc, [
    {'text': 'ד.\t', 'size': 11},
    {'text': 'הוראות מיוחדות', 'underline': True, 'size': 11},
    {'text': ': אין.', 'size': 11},
], space_after=18, indent_right=0.5)

add_p(doc, "יש לכתוב במחברת הבחינה בלבד. יש לרשום \"טיוטה\" בראש כל עמוד המשמש טיוטה.",
      align='center', size=10, space_after=2)
add_p(doc, "כתיבת טיוטה בדפים שאינם במחברת הבחינה עלולה לגרום לפסילת הבחינה.",
      align='center', size=10, space_after=10)

add_p(doc, "השאלות בשאלון זה מנוסחות בלשון רבים, אף על פי כן על כל תלמידה וכל תלמיד להשיב עליהן באופן אישי.",
      bold=True, align='center', size=11, space_after=10)

add_p(doc, "בהצלחה!", bold=True, align='center', size=14, space_after=6)
add_horizontal_continue(doc)

add_page_break(doc)

# ===================== Page header for subsequent pages =====================
# We'll use running text at top of each content page

def page_top(doc, page_num_text=''):
    # a thin separator line style header
    p = add_p(doc, "תולדות עם ישראל, מתכונת אייר תשפ\"ו, מס' 30281",
              size=9, align='right', space_after=6)

page_top(doc)

# ===================== פרק ראשון =====================

add_p(doc, "השאלות", bold=True, size=16, align='center', space_before=6, space_after=8)
add_runs(doc, [
    {'text': 'פרק ראשון', 'bold': True, 'size': 14},
    {'text': '  (29 נקודות)', 'size': 12},
], align='center', space_after=6)
add_p(doc, "ענו על אחת מן השאלות 1–2. בשאלה שבחרתם ענו על כל הסעיפים (29 נקודות).",
      size=11, align='right', space_after=10)

# --------- שאלה 1: המצב הרוחני במושבות ומסע הרבנים (מתשפ"ג) ---------
add_runs(doc, [
    {'text': '1.\t', 'bold': True, 'size': 12},
    {'text': 'המצב הרוחני במושבות ומסע הרבנים', 'bold': True, 'underline': True, 'size': 12},
], space_before=6, space_after=6)

add_p(doc, "לפניכם שלושה קטעים העוסקים במסע הרבנים למושבות (בשנת תרס\"ד). קראו את הקטעים וענו על הסעיפים שאחריהם.",
      size=11, align='justify', space_after=8)

add_p(doc, "קטע א", underline=True, size=11, align='center', space_after=4)
add_p(doc,
      "רוב העולים בעליות הראשונות היו שומרי מסורת, אך הם עמדו אין אונים נגד רוח הכפירה וההתכחשות ששררה בין המורים "
      "והמחנכים של אותה תקופה. הקבוצה האינטלקטואלית שהנחתה והדריכה בהשראתה ובכוח השכלתה את היישוב החדש, "
      "הרסה את הגשרים עם בית אבא ברוסיה או בפולניה, וסבלה מתסביך בכל מה שקשור עם הדת ובכל מה שריח של מסורת "
      "נדף ממנו. קבוצות כחש אלו מצאו גיבוי ומשענת איתנה בפקידי הברון שהיו חניכי התרבות האירופית המתנכרת. בתי הספר "
      "במושבות הברון באותה עת היו מעורטלים מכל ריח של תורה וממושג כלשהו על הדת או המסורת, ואילו האיכרים העניים, "
      "שהיו נתונים לחסדם של פקידי הברון, נאלצו להיכנע ולשלוח את ילדיהם לבתי ספר אלה.",
      size=10, align='justify', space_after=4)
add_p(doc, "(מעובד על פי ש\"ז זוננפלד, האיש על החומה, עמ' 258)",
      size=9, align='left', space_after=8)

add_p(doc, "קטע ב", underline=True, size=11, align='center', space_after=4)
add_p(doc,
      "בכ\"ח באייר תרס\"ד הגיע הרב קוק ליפו, וכבר לאחר כמה שבועות הוא קיבל על עצמו להוביל מסע רבנים אל המתיישבים "
      "בקולוניות [מושבות] ראשון לציון, נס ציונה, רחובות, עקרון וגדרה.",
      size=10, align='justify', space_after=4)
add_p(doc,
      "המטרה של מסע הרבנים למושבות החדשות בארץ־ישראל, שאת חלקן ייסדו חלוצי העלייה השנייה שלא היו שומרי תורה "
      "ומצוות, הייתה לעורר לקירוב לבבות בין היישוב הישן ליישוב החדש, ולקרב את החלוצים עד כמה שאפשר למסורת ישראל "
      "סבא, בנושאי חינוך הילדים, כשרות, שמירת השבת ועוד, עקב מצבם הדתי הירוד של החלוצים.",
      size=10, align='justify', space_after=4)
add_p(doc,
      "חורף כבד בעמק יזרעאל. חושך, ענן וערפל. ארובות השמיים נפתחו. הסוסים מושכים בכבדות את פרסותיהם מתוך הבוץ "
      "הטובעני, והעגלה שוקעת לאיטה. הרב בנימין הורוביץ מחשב בראשו כמה זמן יידרש להם לשוב אל התחנה הקודמת, עד "
      "יעבור זעם; הוא מתלחש עם הרב יעקב משה חרל\"פ, המחריש, משתאה. עיניו בעיני רבו הגדול, הרושפות אש, אור חוזר "
      "מהבזקי הברקים. \"אות משמיים\", חולפת מחשבה מהוססת בליבו של הרב יוסף חיים זוננפלד, והוא שולח מבט ארוך אל הרב "
      "קוק, ממובילי המסע. כל הירידה הזאת אל השפלה ואל העמק, כל היציאה מקרתא דשופרייא, עיר האלוקים – קשה הייתה "
      "לו מלכתחילה. [...] לא בקלות נענה להזמנת הרב מיפו לצאת למסע המושבות, שיסודתו בהררי קודש: לנסות להציל את "
      "מורשת ישראל סבא, שלא תיכרת כליל מפיהם ומאורחות חייהם של החלוצים.",
      size=10, align='justify', space_after=4)
add_p(doc,
      "ההתבטאויות שהוטחו בפניהם פומבית, לא פעם ולא פעמיים, היו לעיתים בוטות ומתריסות. אחדים ממנהיגי המושבות, "
      "ובדרך כלל דווקא הדמויות האינטלקטואליות יותר, התעמתו עם הרבנים בלא כחל וסרק, בחוסר מידה מינימלית של נימוסי "
      "דרך ארץ. אף זאת צריך לזכור שהדבר לא הרתיע את חברי המשלחת מהמשך דרכם; יתרה מזאת, הם לא טמנו ידם בצלחת "
      "ובהחלט השיבו מנה אחת אפיים, דבר דבור על אופניו, גם אם \"בנחת רוח, בשפה ברורה ובנעימה\". כל שאלה זכתה לתשובה, "
      "כל תובנה – גם הביקורתית ביותר – זכתה להתייחסות ולמענה.",
      size=10, align='justify', space_after=4)
add_p(doc, "(מעובד על פי מגד ירחים, 171, כסלו תשע\"ד)",
      size=9, align='left', space_after=8)

add_p(doc, "קטע ג", underline=True, size=11, align='center', space_after=4)
add_p(doc,
      "איני יכול להעלות על הכתב את גודל ההשפעה שהשפיעה נסיעת רבנים זו בכל ארץ־ישראל, ומה אספר כי בכל מקום בואנו "
      "היו עושים כעין יום טוב עד שהיו כולם בטלים ממלאכתם והיו אומרים מי ייתן והיה לנו ביקור רבנים כגון זה לכל הפחות "
      "פעם אחת בשנה, כי ראיית פני הרבנים בלבד עשתה עליהם רושם בל יימחה.",
      size=10, align='justify', space_after=4)
add_p(doc,
      "ועוד שנתגבר בזה כוחם של חרדי המושבות וירדה השפעתם של החופשיים בראותם כבוד מלכים זה שחלקו לרבנים, כי "
      "עד אז היו רואים רק איך שמכבדים למנהלים ודוקטורים שהיו באים אליהם לעיתים, וגם שהיו מלאים התפעלות מדרשות "
      "הרבנים אשר לא יסולאו בפז נגד החינוך החופשי וחילול שבת ובעניין טהרת המשפחה והפרשת תרומות ומעשרות וכו'. "
      "אשרי עין ראתה כל אלה!",
      size=10, align='justify', space_after=4)
add_p(doc, "(מעובד על פי הרב ב\"צ יאדלר, בטוב ירושלים, נצח, תשכ\"ז, עמ' קסו)",
      size=9, align='left', space_after=10)

# Sub-items (29pts = 9+11+9)
add_runs(doc, [
    {'text': 'א.\t', 'bold': True, 'size': 11},
    {'text': 'על פי מה שלמדתם', 'underline': True, 'size': 11},
    {'text': ', הציגו ', 'size': 11},
    {'text': 'שניים', 'underline': True, 'size': 11},
    {'text': ' מן הגורמים למצב הרוחני במושבות שמתואר בקטע א.  ', 'size': 11},
    {'text': '(9 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ב.\t', 'bold': True, 'size': 11},
    {'text': 'על פי הקטעים', 'underline': True, 'size': 11},
    {'text': ', הציגו ', 'size': 11},
    {'text': 'שניים', 'underline': True, 'size': 11},
    {'text': ' מן הקשיים שהיו לרבנים במסעם, והסבירו כיצד התמודדו הרבנים עם ', 'size': 11},
    {'text': 'אחד', 'underline': True, 'size': 11},
    {'text': ' מן הקשיים האלה.  ', 'size': 11},
    {'text': '(11 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ג.\t', 'bold': True, 'size': 11},
    {'text': 'על פי הקטעים ועל פי מה שלמדתם', 'underline': True, 'size': 11},
    {'text': ', כתבו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' תוצאות של "מסע הרבנים".  ', 'size': 11},
    {'text': '(9 נקודות)', 'size': 11},
], space_after=10, indent_right=0.5)

add_page_break(doc)
page_top(doc)

# --------- שאלה 2: גזרת השחיטה (מתשפ"ה שאלה 1) ---------
add_runs(doc, [
    {'text': '2.\t', 'bold': True, 'size': 12},
    {'text': 'ההתמודדות של יהודי גרמניה עם גזרות הנאצים בשנות השלושים', 'bold': True, 'underline': True, 'size': 12},
], space_before=6, space_after=6)

add_p(doc,
      "לפניכם קטע העוסק בגזרת השחיטה שגזרו הנאצים על היהודים בגרמניה בשנות השלושים, "
      "מתוך דבריו של הרב ויינברג בספרו \"שרידי אש\". קראו את הקטע, וענו על השאלות שאחריו.",
      size=11, align='justify', space_after=8)

add_p(doc,
      "\"המלכות הרשעה של הנאצים גזרה איסור על השחיטה בלא הימום הבהמה [שימוש במכת חשמל כדי שהבהמה תאבד הכרה] "
      "קודם שחיטתה. לאחר מכן שללו הרשעים בתחבולות חוקיות אכזריות את האפשרות להביא בשר מן החוץ... נשקפה סכנה "
      "חמורה שמרבית היהודים לא יעמדו בניסיון, חס וחלילה, וייכשלו באכילת נבלות וטרפות.",
      size=10, align='justify', space_after=4)
add_p(doc,
      "בינתיים נשמע הקול כי שלטון הקהילה העברית [היהודית] הגדולה בברלין חושב מחשבות להנהיג במוסדות של הקהילה – "
      "בית החולים, בתי מושב זקנים ועוד – אכילת בשר נבלה וטרפה הנקנה במקולים [איטליזים] של נוכרים. שלטון הקהילה – "
      "שאף על פי שהורכב ברובו הגדול מאנשים ליברליים [רפורמים], נזהר תמיד בנוגע למוסדות הקהילה הרשמיים בענייני "
      "כשרות על פי חוקי ה'שולחן ערוך' והיה סר למשמעתם ולהשגחתם של הרבנים החרדים – טען שבשעת גזרה קשה זו הם "
      "אנוסים בעל כורחם להתיר אכילת נבלות וטרפות... שמועה זו פגעה בנפשנו.",
      size=10, align='justify', space_after=4)
add_p(doc,
      "בצר לנו גמרנו להימלך [החלטנו להתייעץ] עם גדולי ליטא ופולין אם יש להתיר בשעת סכנה זו את ההימום כפי דרישת "
      "הממשלה הרשעה. נסעתי לווילנה, לוורשה וללובלין לשם התייעצות בשאלה חמורה זו הנוגעת לנפש היהדות בגרמניה. "
      "הגאונים הטילו עליי לחבר חיבור בבירור השאלה מהצד ההלכתי... אומנם הראיתי בה פנים [שיקולים] להיתר, אבל ידעתי "
      "מראש כי גאוני ליטא ופולין וגדולי ישראל, מנהיגי היהדות החרדית, לא יסכימו לעולם להיתר שיש בו משום שינוי אופן "
      "השחיטה הנהוג בישראל מדורי דורות, ואף אני בעצמי ליבי נקפני [הציק לי] מאוד לגשת לעניין חמור זה של כשרות השחיטה, "
      "שהוא יסוד היסודות בחיי ישראל...",
      size=10, align='justify', space_after=4)
add_p(doc,
      "כמה פעמים אמרתי לרב של 'עדת ישראל' בברלין, שעמד בראש הלשכה לענייני השחיטה "
      "בגרמניה, שאין לנו לחפש היתרים לשינוי אופן השחיטה... הצורר המזוהם, ראש השלטון הנאצי, ואולי אלפי אלפים שכמותו "
      "יאבדו מן העולם, ודתנו הקדושה תעמוד לעד. על היהודים בגרמניה לעמוד בניסיון למען דתנו הקדושה ולמען אחינו שבכל "
      "הארצות. ואם חס וחלילה נורה להקל בשחיטה זו, עתידים אנו לסכן את השחיטה היהודית בכל העולם. עלינו להראות לכל "
      "העולם שמוכנים אנו למסור את נפשנו על קודשי ישראל, ויראו שונאינו שבאיסור השחיטה לא יעבירו את ישראל על דתו...",
      size=10, align='justify', space_after=4)
add_p(doc,
      "החיבור [שכתבתי] על הימום הבהמות נשלח לכל גדולי ישראל לשם קבלת חוות דעתם. רובם ככולם אסרו את ההימום, ואף "
      "המעטים שנטו להקל בשעת מצוקה ומשום ספק פיקוח נפש, חזרו אחר כך מדעתם הראשונה והצטרפו אל האוסרים, ונמנו "
      "וגמרו שלא להתיר בשום אופן [את ההימום לפני השחיטה]. קיום דת ישראל וקיום היהדות והצלת כבודה דורשים מאיתנו "
      "להביא את הקורבן הקדוש של איסור בשר, ובזכות זה נזכה לגאולה שלמה ולפדות נפשנו\".",
      size=10, align='justify', space_after=4)
add_p(doc, "(מעובד על פי הרב י\"י ויינברג, שרידי אש, חלק ב, סימן ד, כז–כט, ירושלים, תשע\"ב)",
      size=9, align='left', space_after=10)

# Sub-items (29pts = 11+7+11)
add_runs(doc, [
    {'text': 'א.\t', 'bold': True, 'size': 11},
    {'text': 'על פי הקטע', 'underline': True, 'size': 11},
    {'text': ', מה הייתה גזרת השחיטה שגזרו הנאצים בשנות השלושים?', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'על פי מה שלמדתם', 'underline': True, 'size': 11},
    {'text': ', כתבו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' גזרות נוספות שגזרו הנאצים על היהודים בשנות השלושים.  ', 'size': 11},
    {'text': '(11 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ב.\t', 'bold': True, 'size': 11},
    {'text': 'על פי הקטע ועל פי מה שלמדתם על המאפיינים הדתיים של רוב יהודי גרמניה', 'underline': True, 'size': 11},
    {'text': ', הסבירו מדוע חששו הרבנים כי רבים מיהודי גרמניה יאכלו טרפות ונבלות בעקבות גזרת השחיטה.  ', 'size': 11},
    {'text': '(7 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ג.\t', 'bold': True, 'size': 11},
    {'text': 'על פי הקטע', 'underline': True, 'size': 11},
    {'text': ', כתבו כיצד התמודדו רבני גרמניה עם איסור השחיטה. בתשובתכם התייחסו גם לפסיקתם הסופית.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'על פי מה שלמדתם', 'underline': True, 'size': 11},
    {'text': ', הציגו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' דרכי התמודדות אחרות של יהודי גרמניה עם רדיפות הנאצים בשנות השלושים.  ', 'size': 11},
    {'text': '(11 נקודות)', 'size': 11},
], space_after=10, indent_right=0.5)

add_page_break(doc)
page_top(doc)

# ===================== פרק שני =====================

add_runs(doc, [
    {'text': 'פרק שני', 'bold': True, 'size': 14},
    {'text': '  (50 נקודות)', 'size': 12},
], align='center', space_after=6)
add_p(doc, "ענו על שתיים מן השאלות 3–6  (לכל שאלה – 25 נקודות).",
      size=11, align='right', space_after=10)

# --------- שאלה 3: היציאה מחוץ לחומות (מתשפ"ה שאלה 5) ---------
add_runs(doc, [
    {'text': '3.\t', 'bold': True, 'size': 12},
    {'text': 'היציאה מחוץ לחומות', 'bold': True, 'underline': True, 'size': 12},
], space_before=4, space_after=6)

add_runs(doc, [
    {'text': 'א.\t', 'bold': True, 'size': 11},
    {'text': 'כתבו ', 'size': 11},
    {'text': 'שלושה', 'underline': True, 'size': 11},
    {'text': ' גורמים לבניית שכונות חדשות מחוץ לחומות ירושלים.  ', 'size': 11},
    {'text': '(9 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ב.\t', 'bold': True, 'size': 11},
    {'text': 'כתבו על שכונה ', 'size': 11},
    {'text': 'אחת', 'underline': True, 'size': 11},
    {'text': ' שהוקמה בירושלים ביוזמת משה מונטיפיורי.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'הציגו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' פעולות נוספות שעשה מונטיפיורי למען ', 'size': 11},
    {'text': 'יהודי העיר', 'underline': True, 'size': 11},
    {'text': ', וציינו פעולה ', 'size': 11},
    {'text': 'אחת', 'underline': True, 'size': 11},
    {'text': ' שלו למען ', 'size': 11},
    {'text': 'יהודי העולם', 'underline': True, 'size': 11},
    {'text': '  (סך הכול – ', 'size': 11},
    {'text': 'שלוש', 'underline': True, 'size': 11},
    {'text': ' פעולות).  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ג.\t', 'bold': True, 'size': 11},
    {'text': 'בני היישוב הישן יצאו להקים מושבות בארץ־ישראל.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'כתבו על ', 'size': 11},
    {'text': 'שני', 'underline': True, 'size': 11},
    {'text': ' מכשולים שעמדו בדרכם.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'ציינו שמות של ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' מושבות שהוקמו בתקופה זו.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=10, indent_right=0.5)

# --------- שאלה 4: מאורעות תרפ"ט (מתשפ"ד שאלה 6) ---------
add_runs(doc, [
    {'text': '4.\t', 'bold': True, 'size': 12},
    {'text': 'מאורעות תרפ"ט', 'bold': True, 'underline': True, 'size': 12},
], space_before=4, space_after=6)

add_runs(doc, [
    {'text': 'א.\t', 'bold': True, 'size': 11},
    {'text': 'תארו את הרקע לפרוץ פרעות תרפ"ט. בתשובתכם התייחסו למצב ', 'size': 11},
    {'text': 'הכלכלי', 'underline': True, 'size': 11},
    {'text': ' וה', 'size': 11},
    {'text': 'מדיני', 'underline': True, 'size': 11},
    {'text': ' באותה תקופה.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'הציגו את הטיעון שהעלה המופתי חאג\' אמין אל־חוסייני במטרה להתסיס את ההמונים נגד היהודים.  ', 'size': 11},
    {'text': '(9 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ב.\t', 'bold': True, 'size': 11},
    {'text': 'ציינו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' ועדות חקירה בריטיות שהוקמו לאחר פרעות תרפ"ט, ופרטו את המסקנות של ', 'size': 11},
    {'text': 'כל אחת', 'underline': True, 'size': 11},
    {'text': ' מהן.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ג.\t', 'bold': True, 'size': 11},
    {'text': '"הספר הלבן" השני פורסם בשנת 1930. ציינו מי פרסם אותו.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'כתבו על ', 'size': 11},
    {'text': 'שתיים', 'underline': True, 'size': 11},
    {'text': ' מן ההחלטות בספר.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'כתבו מה הייתה המשמעות של "הספר הלבן" השני בנוגע להתחייבות של הבריטים כלפי היישוב היהודי שניתנה '
             'בסוף מלחמת העולם הראשונה ', 'size': 11},
    {'text': 'או', 'underline': True, 'size': 11},
    {'text': ' בנוגע להתחייבות שניתנה בראשית תקופת המנדט.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=10, indent_right=0.5)

add_page_break(doc)
page_top(doc)

# --------- שאלה 5: ארץ-ישראל במלחמת העולם הראשונה (מתשפ"ד שאלה 5) ---------
add_runs(doc, [
    {'text': '5.\t', 'bold': True, 'size': 12},
    {'text': 'ארץ־ישראל בתקופת מלחמת העולם הראשונה', 'bold': True, 'underline': True, 'size': 12},
], space_before=4, space_after=6)

add_runs(doc, [
    {'text': 'א.\t', 'bold': True, 'size': 11},
    {'text': 'לפניכם רשימת מונחים הנוגעים לסוף שלטון העות\'מאנים בארץ־ישראל: "הפתקה האדומה", ביטול '
             '"חוק הקפיטולציות", "מצנפות ירוקות". בחרו ', 'size': 11},
    {'text': 'שניים', 'underline': True, 'size': 11},
    {'text': ' מהם, וכתבו כיצד ', 'size': 11},
    {'text': 'כל אחד', 'underline': True, 'size': 11},
    {'text': ' מן המונחים שבחרתם מבטא את הקשיים של יהודי ארץ־ישראל בימי מלחמת העולם הראשונה. הציגו קושי ', 'size': 11},
    {'text': 'נוסף', 'underline': True, 'size': 11},
    {'text': ' של היהודים בארץ־ישראל בתקופה זו.  ', 'size': 11},
    {'text': '(9 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ב.\t', 'bold': True, 'size': 11},
    {'text': 'הציגו דוגמה להשתתפות יהודים במאמץ המלחמתי של הבריטים נגד העות\'מאנים. בתשובתכם כתבו נימוק ', 'size': 11},
    {'text': 'אחד', 'underline': True, 'size': 11},
    {'text': ' בעד הסיוע לבריטים ונימוק ', 'size': 11},
    {'text': 'אחר', 'underline': True, 'size': 11},
    {'text': ' נגדו.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ג.\t', 'bold': True, 'size': 11},
    {'text': 'כתבו את התוכן של "הצהרת בלפור". כתבו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' תגובות של גדולי ישראל על ההצהרה.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=10, indent_right=0.5)

# --------- שאלה 6: כלכלת ישראל (מתשפ"ג — נשאר) ---------
add_runs(doc, [
    {'text': '6.\t', 'bold': True, 'size': 12},
    {'text': 'כלכלת ישראל בשנים הראשונות להקמת המדינה', 'bold': True, 'underline': True, 'size': 12},
], space_before=4, space_after=6)

add_runs(doc, [
    {'text': 'א.\t', 'bold': True, 'size': 11},
    {'text': 'הציגו ', 'size': 11},
    {'text': 'שלושה', 'underline': True, 'size': 11},
    {'text': ' גורמים לשפל הכלכלי שמדינת ישראל הייתה נתונה בו בשנותיה הראשונות.  ', 'size': 11},
    {'text': '(9 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ב.\t', 'bold': True, 'size': 11},
    {'text': 'הסבירו מהו "משטר הצנע", והציגו ', 'size': 11},
    {'text': 'שני', 'underline': True, 'size': 11},
    {'text': ' קשיים שהוא גרם לתושבי המדינה.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=6, indent_right=0.5)

add_runs(doc, [
    {'text': 'ג.\t', 'bold': True, 'size': 11},
    {'text': '(1) ציינו ', 'size': 11},
    {'text': 'שלושה', 'underline': True, 'size': 11},
    {'text': ' פתרונות נוספים, מלבד משטר הצנע, שסייעו להתמודדות עם המצב הכלכלי הקשה.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': '(2) הסבירו כיצד ', 'size': 11},
    {'text': 'אחד', 'underline': True, 'size': 11},
    {'text': ' מן הפתרונות שציינתם סייע להתמודדות עם המצב הכלכלי הקשה.  ', 'size': 11},
    {'text': '(8 נקודות)', 'size': 11},
], space_after=10, indent_right=0.5)

add_page_break(doc)
page_top(doc)

# ===================== פרק שלישי =====================

add_runs(doc, [
    {'text': 'פרק שלישי', 'bold': True, 'size': 14},
    {'text': '  (21 נקודות)', 'size': 12},
], align='center', space_after=6)
add_p(doc, "ענו על שלוש מן השאלות 7–11  (לכל שאלה – 7 נקודות).",
      size=11, align='right', space_after=10)

# שאלה 7: הגאון מווילנה (מתשפ"ג — נשאר)
add_runs(doc, [
    {'text': '7.\t', 'bold': True, 'size': 11},
    {'text': 'ציינו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' הוראות של הגאון מווילנה בענייני הלכה ומנהג, והציגו את פועלו של ', 'size': 11},
    {'text': 'אחד', 'underline': True, 'size': 11},
    {'text': ' מתלמידיו.', 'size': 11},
], space_after=10, indent_right=0.5)

# שאלה 8: ה"בן איש חי" (מתשפ"ג — נשאר)
add_runs(doc, [
    {'text': '8.\t', 'bold': True, 'size': 11},
    {'text': 'הסבירו מי היה ה"בן איש חי", מדוע זכה לכינוי זה, היכן פעל, ומה היה תפקידו.', 'size': 11},
], space_after=10, indent_right=0.5)

# שאלה 9: הבעש"ט (מתשפ"ה שאלה 7)
add_runs(doc, [
    {'text': '9.\t', 'bold': True, 'size': 11},
    {'text': 'הציגו ', 'size': 11},
    {'text': 'שלוש', 'underline': True, 'size': 11},
    {'text': ' מפעולותיו של הבעש"ט, וציינו את שמותיהם של ', 'size': 11},
    {'text': 'שניים', 'underline': True, 'size': 11},
    {'text': ' מממשיכי דרכו בעולם החסידות.', 'size': 11},
], space_after=10, indent_right=0.5)

# שאלה 10: שאלה 7 מתשפ"ד ללא סעיף ג' — התלמיד בוחר אחד מתוך 4
add_runs(doc, [
    {'text': '10.\t', 'bold': True, 'size': 11},
    {'text': 'בחרו ', 'size': 11},
    {'text': 'אחד', 'underline': True, 'size': 11},
    {'text': ' מן הסעיפים א–ד שלפניכם וענו עליו:', 'size': 11},
], space_after=4, indent_right=0.5)

add_runs(doc, [
    {'text': '\tא.\t', 'bold': True, 'size': 11},
    {'text': 'הסבירו מה הייתה "הפרדת הקהילות" בהונגריה. כתבו על הקשר בינה ובין המאבק ברפורמה.', 'size': 11},
], space_after=4, indent_right=0.5)

add_runs(doc, [
    {'text': '\tב.\t', 'bold': True, 'size': 11},
    {'text': 'בישיבת וולוז\'ין היו כמה שינויים בהשוואה לישיבות שקדמו לה. כתבו ', 'size': 11},
    {'text': 'שני', 'underline': True, 'size': 11},
    {'text': ' שינויים.', 'size': 11},
], space_after=4, indent_right=0.5)

add_runs(doc, [
    {'text': '\tג.\t', 'bold': True, 'size': 11},
    {'text': 'הציגו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' מטרות בהקמת "אגודת ישראל", וציינו ', 'size': 11},
    {'text': 'שני', 'underline': True, 'size': 11},
    {'text': ' אישים שנמנו עם מייסדי התנועה.', 'size': 11},
], space_after=4, indent_right=0.5)

add_runs(doc, [
    {'text': '\tד.\t', 'bold': True, 'size': 11},
    {'text': 'כתבו ', 'size': 11},
    {'text': 'שתי', 'underline': True, 'size': 11},
    {'text': ' דרכים שנקטו עסקני מפלגת השלטון כדי להעביר את ילדי העולים על דתם בשנים הראשונות לאחר '
             'קום המדינה. ציינו מדוע ניסיונות אלה כוונו בעיקר כלפי העולים מארצות האסלאם.', 'size': 11},
], space_after=10, indent_right=0.5)

# שאלה 11: משפט אייכמן (מתשפ"ב שאלה 12)
add_runs(doc, [
    {'text': '11.\t', 'bold': True, 'size': 11},
    {'text': 'הציגו מה היה משפט אייכמן.', 'size': 11},
], space_after=3, indent_right=0.5)
add_runs(doc, [
    {'text': '\t', 'size': 11},
    {'text': 'כתבו מה הייתה ההשפעה של המשפט על החברה הישראלית.', 'size': 11},
], space_after=12, indent_right=0.5)

# Closing
add_p(doc, "בהצלחה!", bold=True, align='center', size=14, space_before=20, space_after=4)
add_p(doc, "זכות היוצרים שמורה למדינת ישראל", size=9, align='center', space_after=2)
add_p(doc, "אין להעתיק או לפרסם אלא ברשות משרד החינוך", size=9, align='center')

doc.save(OUT_DOCX)
print(f"Saved: {OUT_DOCX}")
